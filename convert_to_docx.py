#!/usr/bin/env python3
"""Convert enoca AI Otomasyon MD documentation to DOCX format."""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def read_markdown(md_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        return f.read()

def create_docx(md_content, output_file):
    doc = Document()
    
    # Set document styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(15, 23, 42)
    
    # Heading 1
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(59, 130, 246)
    
    # Heading 2
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(20, 184, 166)
    
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    code_language = ""
    
    while i < len(lines):
        line = lines[i]
        
        # Skip HTML comments and meta lines
        if line.strip().startswith('---') and (i == 0 or lines[i-1].strip() == ''):
            i += 1
            continue
        
        # Code block start/end
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line.strip()[3:]
                code_content = []
            else:
                # Add code block as formatted paragraph
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(236, 72, 153)
                p.paragraph_format.left_indent = Inches(0.3)
                in_code_block = False
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        
        # Tables
        elif line.startswith('|'):
            # Collect table rows
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            # Parse and create table
            rows_data = []
            for row in table_lines:
                cells = [c.strip() for c in row.split('|')[1:-1]]
                if not all(c in ['---', ''] for c in cells) and not any(c in ['---'] for c in cells):
                    rows_data.append(cells)
                elif cells == [''] or all(c == '' for c in cells):
                    pass  # Skip separator rows
                else:
                    pass  # Skip separator rows
            
            if rows_data:
                cols = len(rows_data[0])
                table = doc.add_table(rows=len(rows_data), cols=cols)
                table.style = 'Table Grid'
                
                for ri, row_data in enumerate(rows_data):
                    for ci, cell_text in enumerate(row_data):
                        cell = table.rows[ri].cells[ci]
                        cell.text = cell_text
                        if ri == 0:
                            # Header row
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                            cell._element.get_or_add_tcPr()
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:fill'), '1a365d')
                            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            continue
        
        # Bullet points
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            # Check for nested bullets
            j = i + 1
            while j < len(lines) and (lines[j].startswith('  - ') or lines[j].startswith('- ')):
                if lines[j].startswith('  - '):
                    p = doc.add_paragraph(lines[j][4:], style='List Bullet 2')
                elif lines[j].startswith('- '):
                    break
                j += 1
            i = j
            continue
        
        # Numbered lists
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            p.add_run(line[line.index('. ')+2:])
        
        # Horizontal rule
        elif line.strip() == '':
            pass  # Skip empty lines
        
        # Regular paragraphs
        elif line.strip():
            # Check for bold text
            formatted_line = line
            p = doc.add_paragraph()
            
            # Process bold/italic/code markers
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', formatted_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.color.rgb = RGBColor(139, 92, 246)
                else:
                    p.add_run(part)
        
        i += 1
    
    doc.save(output_file)
    print(f"DOCX saved to: {output_file}")

# Add required imports for XML manipulation
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if __name__ == "__main__":
    md_file = "/Users/osmancagrigenc/Downloads/Enoca Projects/ENOCA_AI_Otomasyon_Dokumantasyonu.md"
    docx_file = "/Users/osmancagrigenc/Downloads/Enoca Projects/ENOCA_AI_Otomasyon_Dokumantasyonu.docx"
    
    md_content = read_markdown(md_file)
    create_docx(md_content, docx_file)
