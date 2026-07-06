#!/usr/bin/env python3
"""Convert ENOCA_n8n_Detayli_Kullanici_Dokumantasyonu.md to DOCX with images."""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = "/Users/osmancagrigenc/Downloads/Enoca Projects"
MD_FILE = os.path.join(BASE_DIR, "ENOCA_n8n_Detayli_Kullanici_Dokumantasyonu.md")
DOCX_FILE = os.path.join(BASE_DIR, "ENOCA_n8n_Detayli_Kullanici_Dokumantasyonu.docx")


def read_markdown(md_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        return f.read()


def add_formatted_text(paragraph, text):
    """Add text with bold/italic/inline-code formatting."""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.color.rgb = RGBColor(139, 92, 246)
        else:
            paragraph.add_run(part)


def add_shaded_header_cell(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    cell._element.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '1a365d')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_docx(md_content, output_file):
    doc = Document()

    # Set document styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(15, 23, 42)

    # Heading 1
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(59, 130, 246)

    # Heading 2
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Calibri'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(20, 184, 166)

    # Heading 3
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Calibri'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(15, 23, 42)

    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    in_mermaid_block = False
    mermaid_content = []

    while i < len(lines):
        line = lines[i]

        # Skip YAML frontmatter separators at top
        if line.strip() == '---' and i == 0:
            i += 1
            continue

        # Code / Mermaid block start/end
        if line.strip().startswith('```'):
            lang = line.strip()[3:].lower()
            if not in_code_block and not in_mermaid_block:
                if lang == 'mermaid':
                    in_mermaid_block = True
                    mermaid_content = []
                else:
                    in_code_block = True
                    code_content = []
            else:
                if in_code_block:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_content))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(236, 72, 153)
                    p.paragraph_format.left_indent = Inches(0.3)
                    in_code_block = False
                elif in_mermaid_block:
                    # Mermaid diagrams rendered as monospace text note
                    p = doc.add_paragraph()
                    run = p.add_run('[Diyagram - Mermaid formatında]')
                    run.italic = True
                    run.font.color.rgb = RGBColor(120, 120, 120)
                    for mline in mermaid_content[:15]:
                        doc.add_paragraph(mline, style='List Bullet 2')
                    in_mermaid_block = False
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        if in_mermaid_block:
            mermaid_content.append(line)
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)

        # Images
        elif line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text, img_path = match.groups()
                # Resolve relative path from BASE_DIR
                if not os.path.isabs(img_path):
                    full_path = os.path.normpath(os.path.join(BASE_DIR, img_path))
                else:
                    full_path = img_path
                if os.path.exists(full_path):
                    try:
                        p = doc.add_paragraph()
                        run = p.add_run()
                        run.add_picture(full_path, width=Inches(5.8))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap = doc.add_paragraph()
                        cap_run = cap.add_run(alt_text)
                        cap_run.italic = True
                        cap_run.font.size = Pt(9)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        doc.add_paragraph(f"[Görsel yüklenemedi: {alt_text} - {e}]")
                else:
                    doc.add_paragraph(f"[Görsel bulunamadı: {full_path}]")

        # Tables
        elif line.startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1

            rows_data = []
            for row in table_lines:
                cells = [c.strip() for c in row.split('|')[1:-1]]
                # Skip separator rows
                if all(re.match(r'^:?-+:?$', c) for c in cells):
                    continue
                if cells == [''] or all(c == '' for c in cells):
                    continue
                rows_data.append(cells)

            if rows_data:
                cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=cols)
                table.style = 'Table Grid'

                for ri, row_data in enumerate(rows_data):
                    for ci in range(cols):
                        cell_text = row_data[ci] if ci < len(row_data) else ''
                        cell = table.rows[ri].cells[ci]
                        if ri == 0:
                            add_shaded_header_cell(cell, cell_text)
                        else:
                            cell.text = cell_text
            continue

        # Bullet points
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
            j = i + 1
            while j < len(lines) and (lines[j].startswith('  - ') or lines[j].startswith('    - ')):
                if lines[j].startswith('    - '):
                    doc.add_paragraph(lines[j][6:], style='List Bullet 3')
                else:
                    doc.add_paragraph(lines[j][4:], style='List Bullet 2')
                j += 1
            i = j - 1

        # Numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)

        # Empty lines
        elif line.strip() == '':
            pass

        # Regular paragraphs
        elif line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line)

        i += 1

    doc.save(output_file)
    print(f"DOCX saved to: {output_file}")


if __name__ == "__main__":
    md_content = read_markdown(MD_FILE)
    create_docx(md_content, DOCX_FILE)
